#!/usr/bin/env python3
"""Generate Word (.docx) documents for TCS Agentic AI product documentation."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    shading.append(sh)

def style_table(table, header_color='1F4E79'):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, header_color)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h

def add_table_from_data(doc, headers, rows, header_color='1F4E79'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    style_table(table, header_color)
    return table


# ============================================================================
# DOCUMENT 1: Product Overview
# ============================================================================
def create_product_overview():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('TCS Agentic AI')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('OpenShift Intelligence Platform\nwith AI-Powered Cluster Operations')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x4A, 0x86, 0xC8)

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver.add_run('\n\nProduct Overview Document\nVersion 1.0 | May 2026\n\nTata Consultancy Services')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # Table of contents placeholder
    add_heading_styled(doc, 'Table of Contents', 1)
    doc.add_paragraph('1. Executive Summary')
    doc.add_paragraph('2. The Problem We Solve')
    doc.add_paragraph('3. Reference Architecture')
    doc.add_paragraph('4. Platform Components')
    doc.add_paragraph('5. AI Agent Architecture')
    doc.add_paragraph('6. Supported LLM Providers')
    doc.add_paragraph('7. Enterprise Integrations')
    doc.add_paragraph('8. Technical Specifications')
    doc.add_paragraph('9. Deployment Architecture')
    doc.add_paragraph('10. Deployment Options')
    doc.add_page_break()

    # 1. Executive Summary
    add_heading_styled(doc, '1. Executive Summary', 1)
    doc.add_paragraph(
        'TCS Agentic AI is an enterprise-grade AI-powered operations platform for Red Hat OpenShift '
        'and Kubernetes clusters. It transforms cluster management from a manual, command-line-driven '
        'process into an intelligent, conversational, and automated experience.'
    )
    doc.add_paragraph(
        'Built on the Model Context Protocol (MCP), Agentic AI connects directly to cluster APIs and '
        'combines real-time observability with AI-powered diagnosis, remediation, and predictive '
        'intelligence — all through a unified web dashboard.'
    )

    # 2. The Problem
    add_heading_styled(doc, '2. The Problem We Solve', 1)
    add_table_from_data(doc,
        ['Challenge', 'Impact'],
        [
            ['Engineers spend 30-60 min diagnosing each incident', 'High MTTR, prolonged outages'],
            ['kubectl/oc CLI requires deep expertise', 'Steep learning curve, slow onboarding'],
            ['Monitoring tools show symptoms, not root causes', 'Engineers investigate manually'],
            ['Knowledge lives in people\'s heads', 'Lost when engineers leave'],
            ['L1/L2 support lacks cluster expertise', 'Escalations slow resolution'],
            ['Multiple tools for different tasks', 'Context switching, fragmented workflow'],
            ['Compliance audits require manual evidence', 'Time-consuming, error-prone'],
        ]
    )

    add_heading_styled(doc, 'How TCS Agentic AI Solves It', 2)
    add_table_from_data(doc,
        ['Solution', 'How'],
        [
            ['AI Diagnosis', 'Pod Doctor auto-diagnoses issues in seconds, not minutes'],
            ['Natural Language', 'Ask "why is my pod crashing?" instead of memorizing oc commands'],
            ['Root Cause Analysis', 'AI correlates logs, events, metrics, and node conditions'],
            ['Knowledge Base', 'Resolutions captured and reused — knowledge stays with the org'],
            ['One-Click Remediation', 'AI suggests fixes, Dry Run validates, Apply executes'],
            ['Single Pane of Glass', 'Dashboard + Chat + Alerts + Metrics in one interface'],
            ['Continuous Compliance', 'Automated security scanning with CIS benchmarks'],
        ]
    )

    # 3. Reference Architecture
    doc.add_page_break()
    add_heading_styled(doc, '3. Reference Architecture', 1)
    doc.add_paragraph(
        'TCS Agentic AI follows a layered architecture deployed natively on OpenShift. '
        'The platform sits within the cluster as a first-class citizen, accessing Kubernetes APIs '
        'directly via a service account with appropriate RBAC bindings.'
    )

    add_heading_styled(doc, '3.1 High-Level Architecture', 2)
    arch_layers = [
        ['Users & Clients', 'Web Browser (Dashboard UI), Claude CLI, VS Code / JetBrains IDE Extensions, MCP Client (SDK / stdio)'],
        ['Network Layer', 'TLS Termination (HTTPS/WSS), OpenShift Route (Edge TLS), HAProxy Load Balancer, stdio Direct Transport'],
        ['Middleware', 'Token/SA Authentication, Token Bucket Rate Limiting, Safety Guardrails, Secret Redaction Filter'],
        ['Agentic AI MCP Server', 'Chat Engine (NLU Parser, Context Memory, ITSM Handler, Action Workflow, Fix Executor), '
         'MCP Tools (33 Modules: Cluster Mgmt, Observability, App Lifecycle, Security & Governance, Operations), '
         'Dashboard (6 Views, Architecture, Heatmap, AI Predictions, Command Bar), '
         'Intelligence (Proactive Agent, Learning Engine, Knowledge Base, Predictive Intel)'],
        ['Data Services', 'PostgreSQL StatefulSet (Port 5432) — Chat History, Audit, Knowledge Base; '
         'Redis StatefulSet (Port 6379) — Cache, Sessions, Rate Limits'],
        ['External Services', 'OpenShift API (15+ Resource Types), LLM Providers (OpenAI, Azure, Claude, Ollama), '
         'ServiceNow (ITSM), Ansible Tower (Automation), Prometheus (Metrics & Alerts)'],
        ['Persistent Storage', 'mcp-data PVC (50Gi RWO — Server State), postgres-data PVC (10Gi RWO — DB), '
         'redis-data PVC (5Gi RWO — Cache)'],
    ]
    add_table_from_data(doc, ['Layer', 'Components'], arch_layers)

    add_heading_styled(doc, '3.2 Data Flow: Chat Query Pipeline', 2)
    doc.add_paragraph(
        'When a user submits a query through the chat interface, the request flows through the following pipeline:'
    )
    pipeline_steps = [
        ['1. Input', 'User submits natural language query via browser (POST /api/chat)'],
        ['2. NLU Parser', 'Intent classification and entity extraction from the query text'],
        ['3. Cache Lookup', 'Redis cache checked for previously computed results — cache hit returns immediately via SSE'],
        ['4. Intent Router', 'Routes to one of three handlers based on intent: '
         'Rule-Based (get, list, logs), ITSM Form (change requests, incidents), or LLM Agent (multi-turn + tool calls)'],
        ['5. Backend Calls', 'Handler queries appropriate backends: K8s API, LLM Provider, or ServiceNow'],
        ['6. SSE Response', 'Results streamed back to browser via Server-Sent Events for real-time display'],
        ['7. Persistence', 'Response cached in Redis and conversation stored in PostgreSQL'],
    ]
    add_table_from_data(doc, ['Step', 'Description'], pipeline_steps)

    add_heading_styled(doc, '3.3 Component Summary', 2)
    components = [
        ['Network', 'Protocols', '4', 'HTTPS, SSE, stdio, WebSocket'],
        ['Security', 'Middleware', '5', 'Auth, Rate Limit, Guardrails, Redaction, CORS'],
        ['Chat', 'NLU + Engines', '5', 'Parser, Memory, ITSM, Workflow, Executor'],
        ['Tools', 'MCP Modules', '33', 'Cluster, Pods, Security, Helm, GitOps, Velero...'],
        ['Dashboard', 'Views + Widgets', '6 + 15', 'Dashboard, Chat, Audit, Hub, Intel, Settings'],
        ['Intelligence', 'AI Services', '8', 'Proactive Agent, Learning, KB, Predictions...'],
        ['Endpoints', 'HTTP APIs', '48+', '/api/chat, /api/actions, /api/intelligence/*'],
        ['LLM', 'Providers', '5', 'OpenAI, Azure, Anthropic, Ollama, Built-in'],
        ['External', 'Integrations', '6', 'OpenShift, LLMs, ServiceNow, Ansible, Prometheus, Redis'],
        ['Data', 'Storage', '3', 'PostgreSQL (8+ tables), Redis (cache), PVC (state)'],
    ]
    add_table_from_data(doc, ['Layer', 'Component', 'Count', 'Details'], components)

    # 4. Platform Components
    doc.add_page_break()
    add_heading_styled(doc, '4. Platform Components', 1)

    add_heading_styled(doc, '4.1 AI-Powered Web Dashboard', 2)
    doc.add_paragraph(
        'The command center for cluster operations. Real-time visibility across pods, nodes, operators, '
        'and namespaces — with AI risk predictions updated every 15 seconds via Server-Sent Events.'
    )
    widgets = [
        'Cluster health overview with version, node count, and operator status',
        'Namespace heatmap showing issue distribution across namespaces',
        'AI Risk Predictions with 100% evidence-based accuracy + LLM deep analysis',
        'Security posture score with vulnerability findings',
        'GitOps sync status (ArgoCD integration)',
        'DR readiness score (Velero backup analysis)',
        'Resource optimization with CPU/memory waste detection',
        'Active alerts from Alertmanager with AI triage',
        'Multi-cluster management via Advanced Cluster Management (ACM)',
    ]
    for w in widgets:
        doc.add_paragraph(w, style='List Bullet')

    add_heading_styled(doc, '4.2 AI Chat Interface', 2)
    doc.add_paragraph(
        'Natural language cluster management. Engineers describe what they need in plain English, '
        'and Agentic AI translates it into cluster operations with AI-powered analysis.'
    )
    examples = [
        '"Troubleshoot pod mlflow-server in namespace mlflow"',
        '"Why is my deployment crashing?"',
        '"Scale nginx to 3 replicas in production"',
        '"Show me all OOMKilled pods"',
        '"Raise a change request for memory increase"',
    ]
    for e in examples:
        doc.add_paragraph(e, style='List Bullet')

    add_heading_styled(doc, '4.3 AI Intelligence Command Center', 2)
    doc.add_paragraph(
        'Unified observability hub with real-time streaming, replacing fragmented monitoring tools.'
    )
    sections = [
        ['Live Alert Feed', 'AI-investigated alerts with severity filtering and one-click remediation'],
        ['Risk Predictions', 'Evidence-based (100%) + LLM deep analysis (70-99%) findings'],
        ['Live Cluster Status', 'SSE-powered real-time counters for pods, nodes, operators, events'],
        ['Root Cause Analysis', 'AI-powered causal chain investigation on any finding'],
        ['Predictive Intelligence', 'Trend-based forecasting of future issues'],
        ['Automation Rules', '"When X happens, do Y" natural language triggers'],
        ['Knowledge Base', 'Learned resolutions from past incidents'],
    ]
    add_table_from_data(doc, ['Section', 'Description'], sections)

    add_heading_styled(doc, '4.4 Fix Execution Engine', 2)
    doc.add_paragraph('Safe, controlled cluster modifications with multiple layers of protection:')
    layers = [
        '1. AI Proposal — AI suggests the fix with risk assessment',
        '2. Dry Run — Server-side simulation before actual execution',
        '3. Approval Chains — Multi-level approval for production changes',
        '4. Guardrails — Blocks dangerous operations (delete namespace, CRDs, etc.)',
        '5. Audit Trail — Every action logged with who, what, when',
        '6. Before/After Comparison — Shows resource limits before and after',
        '7. Post-Fix Monitoring — Auto-polls pod status until Running',
    ]
    for l in layers:
        doc.add_paragraph(l, style='List Bullet')

    # 5. Agent Architecture
    doc.add_page_break()
    add_heading_styled(doc, '5. AI Agent Architecture', 1)
    doc.add_paragraph(
        '12 specialized AI agents for domain-specific cluster operations, orchestrated by an MCP Hub '
        'that can connect to multiple MCP servers simultaneously.'
    )
    agents = [
        ['Diagnostics & Healing', 'Pod/deployment troubleshooting, auto-remediation'],
        ['Security & Compliance', 'SCC analysis, RBAC audit, CIS benchmarks'],
        ['Cluster Operations', 'Node management, resource allocation'],
        ['Upgrade Lifecycle', 'Version advisor, preflight checks, upgrade execution'],
        ['Observability', 'Prometheus queries, alerting, metrics analysis'],
        ['CI/CD & GitOps', 'ArgoCD sync, Tekton pipelines'],
        ['Networking & Mesh', 'Service mesh, network policies, connectivity'],
        ['ITSM & Change Management', 'ServiceNow integration, change requests'],
        ['Multi-Cluster (ACM)', 'Federated cluster management'],
        ['Infrastructure & Virtualization', 'KubeVirt VMs, infrastructure operations'],
        ['Workload Management', 'Deployment scaling, resource optimization'],
        ['Proactive Intelligence', 'Continuous scanning, pattern detection'],
    ]
    add_table_from_data(doc, ['Agent', 'Domain'], agents)

    # 6. LLM Providers
    add_heading_styled(doc, '6. Supported LLM Providers', 1)
    doc.add_paragraph('Agentic AI is vendor-neutral — choose the AI provider that fits your organization:')
    providers = [
        ['Anthropic Claude (Opus/Sonnet/Haiku)', 'Cloud', 'Highest accuracy, complex analysis'],
        ['IBM Watsonx', 'Cloud / On-prem', 'Enterprise AI with IBM support'],
        ['Ollama', 'On-premises', 'Air-gapped environments, data sovereignty'],
        ['Azure OpenAI', 'Cloud', 'Microsoft enterprise ecosystem'],
        ['OpenAI (GPT-4)', 'Cloud', 'General purpose'],
        ['Groq', 'Cloud', 'Ultra-fast inference'],
    ]
    add_table_from_data(doc, ['Provider', 'Deployment', 'Best For'], providers)

    # 7. Integrations
    add_heading_styled(doc, '7. Enterprise Integrations', 1)
    integrations = [
        ['Red Hat OpenShift', 'Native OCP integration via service account'],
        ['Prometheus', 'PromQL queries, metric analysis'],
        ['Alertmanager', 'Alert ingestion, triage, and remediation'],
        ['ServiceNow', 'Incident creation, change requests, approval workflows'],
        ['Ansible Automation Platform', 'Playbook execution for complex remediation'],
        ['ArgoCD', 'GitOps sync status, drift detection'],
        ['Velero', 'Backup status, DR readiness analysis'],
        ['PostgreSQL', 'Persistent chat history, knowledge base, audit logs'],
        ['OpenTelemetry', 'Distributed tracing, metrics export'],
        ['ACM', 'Multi-cluster federation and management'],
    ]
    add_table_from_data(doc, ['Integration', 'Purpose'], integrations)

    # 8. Technical Specs
    add_heading_styled(doc, '8. Technical Specifications', 1)
    specs = [
        ['Runtime', 'Node.js 18+'],
        ['Transport', 'HTTP/SSE (dashboard), stdio (MCP clients)'],
        ['Database', 'PostgreSQL (optional — falls back to in-memory)'],
        ['Authentication', 'OpenShift service account / OPENSHIFT_TOKEN'],
        ['Deployment', 'OpenShift pod, Docker container, or standalone'],
        ['Dashboard', 'Single-page application (no framework dependencies)'],
        ['MCP Protocol', 'Full MCP 1.0 support with tools, resources, prompts'],
        ['API', 'RESTful + Server-Sent Events (SSE)'],
    ]
    add_table_from_data(doc, ['Specification', 'Detail'], specs)

    # 9. Deployment Architecture
    doc.add_page_break()
    add_heading_styled(doc, '9. Deployment Architecture', 1)
    doc.add_paragraph(
        'TCS Agentic AI deploys natively on OpenShift as a set of coordinated resources within '
        'a dedicated namespace (openshift-mcp).'
    )

    add_heading_styled(doc, '9.1 OpenShift Deployment Components', 2)
    deploy_arch = [
        ['Ingress Layer', 'DNS (*.apps.cluster) → OpenShift Route (TLS: edge) → Service (ClusterIP:3000)'],
        ['Compute', 'Deployment: agentic-ai-control-plane (quay.io/karuppucs/openshift-mcp-server:latest, Replicas: 1, Port: 3000, '
         'SA: agentic-ai-server with cluster-reader RBAC, Liveness: /healthz, Readiness: /readyz)'],
        ['Stateful Services', 'PostgreSQL StatefulSet (Port 5432) + Redis StatefulSet (Port 6379)'],
        ['Configuration', 'ConfigMap (environment vars), Secret (API keys, DB credentials), '
         'ServiceAccount (agentic-ai-server + ClusterRoleBinding)'],
        ['Network Security', 'NetworkPolicy (allow from router, allow same namespace)'],
        ['Persistent Volumes', 'mcp-data (50Gi RWO), pg-data (10Gi RWO), redis-data (5Gi RWO)'],
    ]
    add_table_from_data(doc, ['Component', 'Details'], deploy_arch)

    add_heading_styled(doc, '9.2 Technology Stack', 2)
    tech_stack = [
        ['Runtime', 'Node.js 20 on Alpine Linux'],
        ['Platform', 'Red Hat OpenShift Container Platform 4.x'],
        ['Database', 'PostgreSQL 15 (persistent)'],
        ['Cache', 'Redis 7 (in-memory)'],
        ['AI/ML', 'OpenAI GPT-4, Azure OpenAI, Anthropic Claude, Ollama'],
        ['Protocol', 'MCP (Model Context Protocol), REST, SSE, WebSocket'],
        ['ITSM', 'ServiceNow (Change Requests, Incidents)'],
        ['Automation', 'Ansible Automation Platform'],
        ['Monitoring', 'Prometheus, AlertManager'],
        ['Container', 'quay.io/karuppucs/openshift-mcp-server'],
    ]
    add_table_from_data(doc, ['Category', 'Technology'], tech_stack)

    # 10. Deployment Options
    add_heading_styled(doc, '10. Deployment Options', 1)
    deploy = [
        ['OpenShift Pod', 'Deploy as a pod with service account — recommended for production'],
        ['Docker/Podman', 'Container image with cluster credentials mounted'],
        ['Standalone', 'Run directly with OPENSHIFT_TOKEN environment variable'],
        ['Helm Chart', 'Production deployment with configurable values'],
    ]
    add_table_from_data(doc, ['Option', 'Description'], deploy)

    path = os.path.join(OUT_DIR, 'TCS-Agentic-AI-Product-Overview.docx')
    doc.save(path)
    print(f'Created: {path}')
    return path


# ============================================================================
# DOCUMENT 2: Product Comparison
# ============================================================================
def create_comparison():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('TCS Agentic AI\nvs\nOfficial OpenShift MCP Server')
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('\nFeature Comparison & Competitive Analysis\nMay 2026')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    add_heading_styled(doc, 'Executive Summary', 1)
    add_table_from_data(doc,
        ['Aspect', 'Official OpenShift MCP Server', 'TCS Agentic AI'],
        [
            ['Purpose', 'CLI tool for AI IDEs to query K8s clusters', 'Full enterprise platform for AI-powered cluster operations'],
            ['Interface', 'No UI — MCP protocol only (used inside IDEs)', 'Web dashboard + AI chat + MCP server + SSE live streaming'],
            ['Target User', 'Developers using Claude Desktop / VS Code', 'SREs, Platform Engineers, Operations Teams, L1-L3 Support'],
            ['Architecture', 'Single Go binary, stateless', 'Node.js platform with PostgreSQL, multi-LLM, SSE, MCP Hub'],
            ['Remediation', 'Read-only / manual approval only', 'AI-recommended fixes with Dry Run + one-click Apply'],
            ['Intelligence', 'None — raw data only', 'AI Risk Predictions, RCA, Predictive Intel, Knowledge Base'],
        ]
    )

    categories = [
        ('Core Kubernetes Operations', [
            ['List/Get/Delete pods', 'Yes', 'Yes'],
            ['Pod logs retrieval', 'Yes', 'Yes + AI-analyzed logs'],
            ['Pod exec (run commands)', 'Yes', 'Yes'],
            ['Resource CRUD (YAML)', 'Yes', 'Yes + guardrails + approvals'],
            ['Scale deployments', 'Yes', 'Yes + impact preview'],
            ['Namespace listing', 'Yes', 'Yes + namespace health heatmap'],
            ['Node logs & stats', 'Yes', 'Yes + real-time metrics via SSE'],
            ['Events listing', 'Yes', 'Yes + AI-correlated events'],
            ['Helm / Tekton / KubeVirt / OSSM', 'Yes', 'Yes'],
        ]),
        ('AI & Intelligence', [
            ['AI Chat Interface', 'No', 'Yes — natural language cluster management'],
            ['Multi-LLM Support', 'No', 'Claude, Watsonx, Ollama, Azure OpenAI, OpenAI, Groq'],
            ['AI Risk Predictions', 'No', 'Yes — 100% evidence-based + LLM deep analysis'],
            ['Root Cause Analysis', 'No', 'Yes — AI-powered causal chain investigation'],
            ['Pod Doctor (auto-diagnosis)', 'No', 'Yes — automated pod health diagnosis'],
            ['Smart OOM Recommendations', 'No', 'Yes — metrics-based memory sizing'],
            ['NLU + Conversation Memory', 'No', 'Yes — context-aware multi-turn conversations'],
        ]),
        ('Dashboard & Visualization', [
            ['Web Dashboard', 'No', 'Yes — full-featured SPA'],
            ['Cluster Health Overview', 'No', 'Yes — real-time health cards'],
            ['Namespace Heatmap', 'No', 'Yes — visual namespace health'],
            ['Security Posture Score', 'No', 'Yes — scored with findings'],
            ['GitOps / DR / Optimization', 'No', 'Yes — ArgoCD, Velero, resource analysis'],
        ]),
        ('Live Streaming & Monitoring', [
            ['SSE Live Cluster Stream', 'No', 'Yes — 15s real-time updates'],
            ['Live Pulse Indicator', 'No', 'Yes — connection status'],
            ['Real-time Event Feed', 'No', 'Yes — warning events streamed'],
        ]),
        ('Remediation & Fix Execution', [
            ['One-click Fix Application', 'No', 'Yes — Apply Fix buttons'],
            ['Dry Run Mode', 'No', 'Yes — preview changes safely'],
            ['Before/After Comparison', 'No', 'Yes — resource limits diff'],
            ['Post-fix Pod Monitoring', 'No', 'Yes — auto-poll until Running'],
            ['Guardrails & Safety Checks', 'No', 'Yes — blocks dangerous operations'],
            ['Approval Chains', 'No', 'Yes — multi-level approval workflows'],
        ]),
        ('Enterprise Integrations', [
            ['ServiceNow (ITSM)', 'No', 'Yes — incidents, change requests'],
            ['Ansible Automation', 'No', 'Yes — playbook execution'],
            ['Prometheus/Alertmanager', 'Partial', 'Yes — full integration + alert triage'],
            ['Multi-Cluster (ACM)', 'No', 'Yes — manage multiple clusters'],
            ['Velero / ArgoCD', 'No', 'Yes — DR + GitOps monitoring'],
        ]),
        ('Knowledge & Learning', [
            ['Knowledge Base', 'No', 'Yes — learned resolutions'],
            ['Automation Rules', 'No', 'Yes — "When X happens, do Y"'],
            ['Playbooks', 'No', 'Yes — team runbooks'],
            ['Persistent Chat History', 'No', 'Yes — PostgreSQL-backed'],
        ]),
        ('Security & Compliance', [
            ['Security Audit / CIS Benchmarks', 'No', 'Yes'],
            ['SCC Advisor / RBAC Analysis', 'No', 'Yes'],
            ['Policy Generation', 'No', 'Yes — network policies, RBAC'],
            ['Data Redaction', 'No', 'Yes — sensitive data masking'],
        ]),
        ('Advanced Operations', [
            ['Upgrade Advisor + Preflight', 'No', 'Yes — risk analysis + validation'],
            ['Emergency Actions', 'No', 'Yes — cordon/drain/force-delete'],
            ['Bulk Operations', 'No', 'Yes — batch fix crashed pods'],
            ['Impact Analysis', 'No', 'Yes — blast radius assessment'],
            ['12 Specialized AI Agents', 'No', 'Yes — domain-specific agents'],
        ]),
    ]

    for cat_name, rows in categories:
        doc.add_page_break()
        add_heading_styled(doc, cat_name, 1)
        add_table_from_data(doc,
            ['Capability', 'Official MCP Server', 'TCS Agentic AI'],
            rows
        )

    # Architecture comparison
    doc.add_page_break()
    add_heading_styled(doc, 'Architecture Comparison', 1)

    add_heading_styled(doc, 'Official OpenShift MCP Server', 2)
    doc.add_paragraph('AI IDE (Claude/VS Code) --> MCP Server (Go binary) --> K8s API Server')
    doc.add_paragraph('Single binary, stateless, IDE-only access. No dashboard, no chat, no persistence.')

    add_heading_styled(doc, 'TCS Agentic AI', 2)
    arch_text = (
        'Web Browser / AI IDE --> TCS Agentic AI Platform --> K8s API Server + Prometheus + '
        'Alertmanager + ServiceNow + ArgoCD\n\n'
        'Platform includes: Dashboard API, Chat API + NLU, MCP Hub + 12 Agents, Pod Doctor, '
        'RCA Engine, Fix Executor, Knowledge Base, Guardrails, LLM Layer (multi-provider), '
        'PostgreSQL persistence, SSE live streaming'
    )
    doc.add_paragraph(arch_text)

    add_heading_styled(doc, 'Summary', 1)
    doc.add_paragraph(
        'The official OpenShift MCP Server is an excellent developer tool — a lightweight, portable '
        'binary that gives AI IDEs read/write access to Kubernetes clusters. '
        'TCS Agentic AI is an enterprise operations platform that goes far beyond MCP protocol '
        'support, providing AI intelligence, auto-remediation, enterprise integrations, and a full '
        'web-based operations center.'
    )

    path = os.path.join(OUT_DIR, 'TCS-Agentic-AI-vs-Official-MCP-Comparison.docx')
    doc.save(path)
    print(f'Created: {path}')
    return path


# ============================================================================
# DOCUMENT 3: Customer Benefits
# ============================================================================
def create_benefits():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('TCS Agentic AI')
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('\nCustomer Benefits & Industry Value Proposition\nMay 2026\n\nTata Consultancy Services')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    # Benefits by stakeholder
    add_heading_styled(doc, 'Benefits by Stakeholder', 1)

    add_heading_styled(doc, 'For Platform Engineering / SRE Teams', 2)
    add_table_from_data(doc,
        ['Benefit', 'How Agentic AI Delivers'],
        [
            ['50-70% MTTR reduction', 'AI diagnoses root cause in seconds. Pod Doctor auto-analyzes container states, logs, events, and metrics'],
            ['Proactive risk detection', '100% evidence-based risk predictions from live cluster APIs. Node saturation, operator degradation detected automatically'],
            ['Eliminate context switching', 'Single dashboard replaces Prometheus + Grafana + kubectl + ServiceNow + Slack'],
            ['Safe remediation', 'Dry Run validates changes. Guardrails block dangerous operations. Approval chains enforce change management'],
            ['24/7 autonomous monitoring', 'Proactive Agent continuously scans. SSE live streaming provides Prometheus-like visibility'],
        ]
    )

    add_heading_styled(doc, 'For L1/L2 Support Teams', 2)
    add_table_from_data(doc,
        ['Benefit', 'How Agentic AI Delivers'],
        [
            ['No kubectl expertise required', 'Natural language: "Show me all crashing pods" works like talking to a senior engineer'],
            ['Guided remediation', 'AI provides step-by-step fix with one-click apply. L1 resolves issues that previously needed L3'],
            ['Built-in knowledge base', 'Past resolutions captured automatically. Similar issues get instant answers'],
            ['Faster onboarding', 'New team members productive in hours, not weeks'],
            ['Reduced escalations', 'AI handles 60-80% of common issues autonomously'],
        ]
    )

    add_heading_styled(doc, 'For Engineering Managers / Directors', 2)
    add_table_from_data(doc,
        ['Benefit', 'How Agentic AI Delivers'],
        [
            ['Operational cost reduction', 'AI automates L1/L2 tasks — skilled engineers focus on architecture'],
            ['Knowledge retention', 'Expertise stays in Knowledge Base when engineers leave'],
            ['Compliance automation', 'Continuous CIS scanning, SCC analysis, RBAC audit — always audit-ready'],
            ['Multi-cluster visibility', 'Single pane of glass across all clusters via ACM integration'],
            ['Vendor-neutral AI', 'Choose Claude, Watsonx, Ollama (on-prem), Azure OpenAI — no lock-in'],
        ]
    )

    add_heading_styled(doc, 'For CTO / CIO / IT Leadership', 2)
    add_table_from_data(doc,
        ['Benefit', 'How Agentic AI Delivers'],
        [
            ['ROI within 3 months', 'MTTR reduction + L1/L2 automation = measurable savings from day one'],
            ['Data sovereignty', 'Ollama/Watsonx on-prem: AI stays within infrastructure. Data redaction masks sensitive info'],
            ['Enterprise compliance', 'RBAC, approval chains, audit trails, ServiceNow integration'],
            ['Scalable operations', 'AI scales with cluster count. More clusters does not mean more engineers'],
            ['Innovation accelerator', 'Teams freed from operational toil focus on business value'],
        ]
    )

    # Industry Alignment
    doc.add_page_break()
    add_heading_styled(doc, 'Industry Alignment', 1)

    add_heading_styled(doc, 'SRE Best Practices (Google SRE Book)', 2)
    add_table_from_data(doc,
        ['SRE Principle', 'Agentic AI Implementation'],
        [
            ['Reduce toil', 'AI automates repetitive diagnosis and remediation'],
            ['Monitor meaningfully', 'Evidence-based alerts, not threshold-based noise'],
            ['Post-mortem culture', 'RCA engine builds causal chains automatically'],
            ['Error budgets', 'Risk predictions quantify reliability impact'],
            ['Automation', 'Automation rules: "When X happens, do Y"'],
        ]
    )

    add_heading_styled(doc, 'ITIL / ITSM Alignment', 2)
    add_table_from_data(doc,
        ['ITIL Process', 'Agentic AI Integration'],
        [
            ['Incident Management', 'AI-powered triage, auto-diagnosis, guided resolution'],
            ['Change Management', 'ServiceNow integration, approval chains, CR tracking'],
            ['Problem Management', 'Root Cause Analysis, Knowledge Base, pattern learning'],
            ['Service Level Management', 'Proactive risk detection prevents SLA breaches'],
            ['Knowledge Management', 'AI-captured resolutions, team playbooks, episodic memory'],
        ]
    )

    add_heading_styled(doc, 'AIOps Market Standards (Gartner / Forrester)', 2)
    add_table_from_data(doc,
        ['AIOps Capability', 'Agentic AI'],
        [
            ['Data ingestion', 'Direct Kubernetes API + Prometheus + Alertmanager + Events'],
            ['Pattern recognition', 'LLM-powered correlation across pods, nodes, events, operators'],
            ['Root cause determination', 'Causal chain analysis with timeline reconstruction'],
            ['Predictive analytics', 'Trend-based capacity and failure forecasting'],
            ['Automated remediation', 'One-click fix with dry run, guardrails, and approval'],
            ['Collaboration', 'Chat-based workflow with persistent conversation history'],
        ]
    )

    # Competitive Analysis
    doc.add_page_break()
    add_heading_styled(doc, 'Competitive Analysis', 1)

    add_heading_styled(doc, 'vs Traditional Monitoring (Datadog, Dynatrace, New Relic)', 2)
    add_table_from_data(doc,
        ['Aspect', 'Traditional Monitoring', 'TCS Agentic AI'],
        [
            ['Pricing', 'Per-host / per-GB licensing', 'Fixed deployment cost — no per-host fees'],
            ['Data residency', 'SaaS — data leaves your environment', 'On-premises option — data stays internal'],
            ['Intelligence', 'Anomaly detection (needs training)', 'AI diagnosis from day one'],
            ['Remediation', 'Alerts only — manual', 'AI-guided fix with one-click apply'],
            ['K8s depth', 'General purpose with plugins', 'Kubernetes-native — built for OCP/K8s'],
        ]
    )

    add_heading_styled(doc, 'vs kubectl / oc CLI', 2)
    add_table_from_data(doc,
        ['Aspect', 'CLI', 'TCS Agentic AI'],
        [
            ['Learning curve', 'Steep — hundreds of commands', 'Natural language — "show crashing pods"'],
            ['Diagnosis', 'Manual correlation', 'Automated across all data sources'],
            ['Remediation', 'Manual YAML editing', 'One-click with dry run preview'],
            ['Audit trail', 'None', 'Built-in with PostgreSQL'],
            ['Knowledge sharing', 'Wiki pages that go stale', 'Live Knowledge Base that grows'],
        ]
    )

    # ROI
    doc.add_page_break()
    add_heading_styled(doc, 'ROI Analysis', 1)
    doc.add_paragraph('Assumptions: 20 clusters, 5000+ pods, 10-person platform team, 50 incidents/week')
    add_table_from_data(doc,
        ['Category', 'Before', 'After', 'Improvement'],
        [
            ['Incident resolution time', '45 min avg', '10 min avg', '78% reduction'],
            ['L1/L2 escalations', '30/week', '8/week', '73% reduction'],
            ['New engineer onboarding', '4 weeks', '1 week', '75% faster'],
            ['Compliance audit prep', '2 weeks/quarter', '2 days/quarter', '80% reduction'],
            ['Monitoring tool licenses', 'Multiple SaaS', 'Single platform', 'Consolidation'],
            ['On-call burden', 'Full manual triage', 'AI-assisted', 'Reduced fatigue'],
        ]
    )

    # Use Cases
    add_heading_styled(doc, 'Use Cases', 1)

    add_heading_styled(doc, 'Use Case 1: OOMKill Resolution', 2)
    doc.add_paragraph('Before: Engineer paged at 2 AM. SSH into cluster. Run multiple oc commands. Analyze logs. Calculate new limits. Edit YAML. Apply. Watch rollout. Total: 35 minutes.')
    doc.add_paragraph('With Agentic AI: Dashboard shows OOMKill with 100% confidence. AI recommends exact memory limit based on actual usage metrics. Dry Run preview. One-click Apply. Auto-poll shows pods Running. Total: 2 minutes.')

    add_heading_styled(doc, 'Use Case 2: CrashLoopBackOff Investigation', 2)
    doc.add_paragraph('Before: Read 500 lines of logs. Check events. Check node conditions. Check Prometheus. Total: 45 minutes.')
    doc.add_paragraph('With Agentic AI: Click "Root Cause". AI correlates pod logs + events + node conditions + metrics. Returns full causal chain with recommended fix. Total: 30 seconds.')

    add_heading_styled(doc, 'Use Case 3: Cluster Upgrade Planning', 2)
    doc.add_paragraph('Before: Read release notes. Check compatibility. Verify node health. Check deprecated APIs. Create CR manually. Total: 2-3 days.')
    doc.add_paragraph('With Agentic AI: Single command: "Plan upgrade to 4.16." AI runs preflight checks, assesses risk, generates change request. Total: 15 minutes.')

    # Implementation Timeline
    add_heading_styled(doc, 'Implementation Timeline', 1)
    add_table_from_data(doc,
        ['Phase', 'Duration', 'Deliverables'],
        [
            ['Phase 1: Deploy', '1 day', 'Agentic AI deployed on OpenShift'],
            ['Phase 2: Configure', '1-2 days', 'LLM provider, ServiceNow, RBAC configured'],
            ['Phase 3: Adopt', '1 week', 'Team onboarded, Knowledge Base populated'],
            ['Phase 4: Optimize', '2-4 weeks', 'Automation rules, playbooks, feedback loop'],
            ['Phase 5: Scale', 'Ongoing', 'Multi-cluster rollout, agent customization'],
        ]
    )

    path = os.path.join(OUT_DIR, 'TCS-Agentic-AI-Customer-Benefits.docx')
    doc.save(path)
    print(f'Created: {path}')
    return path


# ============================================================================
# MAIN
# ============================================================================
if __name__ == '__main__':
    p1 = create_product_overview()
    p2 = create_comparison()
    p3 = create_benefits()
    print(f'\nAll documents created in: {OUT_DIR}')
    print(f'  1. {os.path.basename(p1)}')
    print(f'  2. {os.path.basename(p2)}')
    print(f'  3. {os.path.basename(p3)}')
